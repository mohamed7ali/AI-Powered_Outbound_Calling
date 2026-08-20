"""Local document extraction for organization knowledge ingestion."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    title: str
    mime_type: str
    text: str
    page_count: int | None = None


def extract_text(path: str | Path, *, mime_type: str | None = None) -> ExtractedDocument:
    file_path = Path(path)
    kind = (mime_type or "").lower()
    suffix = file_path.suffix.lower()
    if suffix == ".pdf" or "pdf" in kind:
        return _extract_pdf(file_path)
    if suffix in {".docx", ".doc"} or "word" in kind or "officedocument" in kind:
        if suffix == ".doc":
            raise ValueError("Legacy .doc is not supported; convert it to .docx first")
        return _extract_docx(file_path)
    if suffix in {".txt", ".md", ".csv", ".json"} or kind.startswith("text/"):
        return ExtractedDocument(
            title=file_path.stem,
            mime_type=mime_type or "text/plain",
            text=file_path.read_text(encoding="utf-8", errors="replace"),
        )
    raise ValueError(f"Unsupported document type: {file_path.name}")


def _extract_pdf(path: Path) -> ExtractedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(f"[صفحة {index}]\n{value}" for index, value in enumerate(pages, 1) if value)
    return ExtractedDocument(
        title=path.stem,
        mime_type="application/pdf",
        text=text,
        page_count=len(reader.pages),
    )


def _extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                tables.append(" | ".join(cells))
    text = "\n".join(paragraphs + tables)
    return ExtractedDocument(
        title=path.stem,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        text=text,
    )
